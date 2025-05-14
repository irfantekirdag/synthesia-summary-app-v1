import streamlit as st
import requests
from bs4 import BeautifulSoup
from transformers import pipeline
from fpdf import FPDF
import re
import io
from docx import Document
from docx.shared import Pt, Inches
from datetime import datetime
import qrcode
from deep_translator import GoogleTranslator
from langdetect import detect
from language_tool_python import LanguageTool


lang_map = {
    "Türkçe": None,
    "English": "en",
    "Deutsch (Almanca)": "de",
    "Français": "fr",
    "Español": "es",
    "Italiano": "it"
}

class PDF(FPDF):
    def __init__(self):
        super().__init__()
        self.add_font('DejaVu', '', 'DejaVuSans.ttf', uni=True)
        self.add_font('DejaVu', 'B', 'DejaVuSans-Bold.ttf', uni=True)

    def header(self):
        self.image("synthesia_logo_light.png", x=10, y=8, w=30)
        self.set_xy(10, 25)
        self.set_font("DejaVu", style="B", size=14)
        self.cell(0, 10, PROJECT_TITLE, ln=True, align='L')
        self.ln(15)
        self.line(10, 35, 200, 35)

    def footer(self):
        self.set_y(-15)
        self.set_font("DejaVu", size=8)
        self.cell(0, 10, f"Sayfa {self.page_no()}", align='C')

def fix_sentences(text):
    text = re.sub(r"\s+", " ", text.strip())
    text = re.sub(r'[^\x00-\x7FğüşöçıĞÜŞİÖÇ.,;!?()\-–\[\]\'\"\n ]+', '', text)
    text = re.sub(r"(?<![.!?])\s*\n+\s*", ". ", text)
    text = re.sub(r"\b([a-zA-ZğüşöçıĞÜŞİÖÇ])\s*\.\s*([a-zA-ZğüşöçıĞÜŞİÖÇ])\b", r"\1\2", text)
    sentences = re.split(r'(?<=[.!?]) +', text)
    cleaned = ' '.join(s[0].upper() + s[1:] if len(s) > 1 else s.upper() for s in sentences if s.strip())
    return cleaned

def grammar_correct(text, lang_hint=None):
    try:
        lang = lang_hint if lang_hint else detect(text)
        tool = LanguageTool(public_api=True, language=lang)
        return tool.correct(text)
    except:
        return text

st.set_page_config(page_title="AI Web Özetleyici", page_icon="🧠")

url = st.text_input("📎 Özetlemek istediğiniz web sayfasının URL'sini girin:", key="main_url")

with st.sidebar:
    st.image("synthesia_logo_dark.png", width=140)
    st.markdown("### Synthesia™")
    st.markdown("<p style='color: #cccccc;'>Yapay zekâ destekli özetleyiciye hoş geldiniz!</p>", unsafe_allow_html=True)

if not url:
    st.stop()

st.title(PROJECT_TITLE)
output_format = st.radio("📄 Çıktı formatı seçin:", ["PDF", "TXT", "DOCX"])
lang_choice = st.selectbox("🌍 Hangi dile çevrilsin? (isteğe bağlı)", list(lang_map.keys()))
selected_lang = lang_map[lang_choice]
summary_type = st.selectbox("🧠 Özet türünü seçin:", ["Kısa Özet", "Madde Madde Özet", "Uzun Özet", "Çok Uzun Özet"])

if st.button("📝 Sayfayı Özetle"):
    try:
        response = requests.get(url)
        response.raise_for_status()
        soup = BeautifulSoup(response.text, 'html.parser')
        paragraphs = soup.find_all('p')
        metin = " ".join(p.get_text().strip() for p in paragraphs if len(p.get_text().strip()) > 0)

        if not metin.strip():
            st.warning("Sayfada yeterli yazı içeriği bulunamadı.")
            st.stop()

        title_tag = soup.find('title')
        meta_desc = soup.find('meta', attrs={'name': 'description'})
        title_text = title_tag.get_text().strip() if title_tag else "ozet"
        if meta_desc and meta_desc.get('content'):
            title_text += " - " + meta_desc['content'].strip()
        file_base = re.sub(r'\W+', '_', title_text.lower())[:30]

        summarizer = pipeline("summarization", model="facebook/bart-large-cnn")
        chunks = [metin[i:i+1000] for i in range(0, len(metin), 1000)]
        summaries = []
        progress_text = st.empty()
        progress_bar = st.progress(0)

        for i, chunk in enumerate(chunks):
            if summary_type == "Kısa Özet":
                summary = summarizer(chunk, max_length=60, min_length=30, do_sample=False)[0]['summary_text']

            elif summary_type == "Madde Madde Özet":
                raw_summary = summarizer(chunk, max_length=120, min_length=60, do_sample=False)[0]['summary_text']
                raw_summary = re.sub(r"\s+", " ", raw_summary.strip())
                raw_summary = re.sub(r"(?<=[a-zA-ZğüşöçıĞÜŞİÖÇ])\.(?=[a-zA-ZğüşöçıĞÜŞİÖÇ])", ". ", raw_summary)
                raw_summary = re.sub(r"(?<!\.)\n+", ". ", raw_summary)
                raw_summary = re.sub(r"(\w)-(\w)", r"\1\2", raw_summary)
                raw_summary = re.sub(r"(?<![.!?]) (?=[A-Z])", ". ", raw_summary)
                sentences = raw_summary.split('. ')
                summary = "\n".join([f"• {s.strip()}" for s in sentences if len(s.strip()) > 10])

            elif summary_type == "Çok Uzun Özet":
                summary = summarizer(chunk, max_length=300, min_length=180, do_sample=False)[0]['summary_text']

            elif summary_type == "Uzun Özet":
                summary = summarizer(chunk, max_length=200, min_length=100, do_sample=False)[0]['summary_text']

            else:
                summary = summarizer(chunk, max_length=120, min_length=60, do_sample=False)[0]['summary_text']

            summaries.append(summary)
            progress_bar.progress((i + 1) / len(chunks))
            progress_text.markdown(f"💬 {i + 1}/{len(chunks)} parça özetlendi")

        progress_bar.empty()
        progress_text.empty()

        if summary_type == "Madde Madde Özet":
            final_summary = "\n".join(summaries)
        else:
            final_summary = fix_sentences(" ".join(summaries))

        if selected_lang:
            try:
                title_text = GoogleTranslator(source='auto', target=selected_lang).translate(title_text)
                final_summary = GoogleTranslator(source='auto', target=selected_lang).translate(final_summary)
            except Exception as e:
                st.error(f"❌ Çeviri yapılamadı: {str(e)}")

        if summary_type != "Madde Madde Özet":
            final_summary = grammar_correct(final_summary, lang_hint=selected_lang)

        filename = f"{file_base}_ozet.{output_format.lower()}"
        source_label = "Bu özetin kaynağı:"

        if output_format == "PDF":
            pdf = PDF()
            pdf.add_page()
            pdf.set_font("DejaVu", size=12)
            for line in final_summary.split("\n"):
                pdf.multi_cell(0, 10, line)

            qr = qrcode.make(url)
            qr_path = "qr_temp.png"
            qr.save(qr_path)
            pdf.image(qr_path, x=160, y=pdf.get_y()+10, w=30)
            pdf.ln(35)
            pdf.cell(0, 10, f"{source_label} {url}", ln=True)
            pdf_output = pdf.output(dest='S').encode('latin1')
            pdf_buffer = io.BytesIO()
            pdf_buffer.write(pdf_output)
            pdf_buffer.seek(0)
            st.download_button("📥 Özeti indir (PDF)", pdf_buffer, file_name=filename)

        elif output_format == "TXT":
            buffer = io.BytesIO()
            buffer.write(f"{title_text}\n\n{final_summary}".encode("utf-8"))
            buffer.seek(0)
            st.download_button("📥 Özeti indir (TXT)", buffer, file_name=filename)

        elif output_format == "DOCX":
            doc = Document()
            doc.add_picture("synthesia_logo_light.png", width=Inches(2.5))
            doc.add_heading(PROJECT_TITLE, level=1)
            doc.add_paragraph(f"Kaynak: {title_text}")
            doc.add_paragraph(f"Dil: {lang_choice}   Tarih: {datetime.today().strftime('%d.%m.%Y')}")
            doc.add_paragraph("\n" + "-" * 40 + "\n")
            doc.add_paragraph(final_summary)
            doc.add_paragraph("\n" + source_label + " " + url)
            docx_buffer = io.BytesIO()
            doc.save(docx_buffer)
            docx_buffer.seek(0)
            st.download_button("📥 Özeti indir (DOCX)", docx_buffer, file_name=filename)

        st.success(f"✅ Özet başarıyla oluşturuldu ve indirilmeye hazır: `{filename}`")

    except Exception as e:
        st.error(f"Hata oluştu: {str(e)}")
